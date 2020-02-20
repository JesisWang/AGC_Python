'''
Created on 2019年8月7日

@author: JesisW
'''
import scipy.io as sci
import time
import pandas as pd
import numpy as np
import sys
import os
import matplotlib.pyplot as plt
import FIGUREplot
from AGCanalyse.check_and_deal_data import Check_data,deal_data
from AGCanalyse.Base_analyse import MX,HB,GD,operation_analyse
sys.path.insert(0, r'D:\pyworkspace\Data_Get') # sys.path.append(path)也可以，只不过查询时是在最后位置，加载包的时候是按顺序查找包的
# from getAPI import getStaWork
from docx import Document
from docx.shared import Inches
import time
import datetime
import lxml
from functools import wraps
from win32com import client
'''
    文件取数接C:取文件的时候index是不含有时间的
    path = 'C:/Users/JesisW/Desktop/XFdata.mat'
    data = sci.loadmat(path)
    M = data['XFdata']
    val = M[0,0]
    df = pd.DataFrame(val['data1203'],columns = ['Agc','Pdg','Pall'])
    接口取数接C:取接口的时候，index是含有时间的
    sta_data_date = '2019-05-16 00:00:00'
    end_data_date = '2019-05-16 22:00:00'
    sta_config=pd.read_excel('D:/pyworkspace/Data_Get/assets/sta_cl_ah_config.xlsx')
    sta_codes,sta_names=sta_config['code'],sta_config['name']
    index =      [  0,    1,    2,   5,    6,    7,   14,   15,    16,   33   ,  4]  #调频电站序号
    index_name = ['新丰','云河','准大','海丰','河源','宣化','同达','上都','平朔','鲤鱼江','兴和']
    index_loc =  ['MX' ,'GD',' MX' ,'GD', 'GD', 'HB', 'HB', 'HB', 'HB', 'GD' , 'MX']
    #  i         [  0     1     2    3     4     5     6     7     8     9   ,  10]
    i = 0 # 电站i = 0~9
    Aim_codes,Aim_names,Aim_names_Chinese = sta_codes[index[i]],sta_names[index[i]],index_name[i]
    
    A = getStaWork(sta_data_date,end_data_date,Aim_codes)
    df,df1,df2 = A.getEMSBat()
    Agc1,Agc2 = A.getEMSAgc()
    DF = pd.merge(df,Agc1,right_index=True,left_index=True,how='inner')
    time = pd.DataFrame(DF.index,columns =['time'])
    time.columns = ['time']
    data = pd.DataFrame(columns=['Agc','Pdg','Pall','Pbat'])
    data['Agc'],data['Pdg'],data['Pbat'] = DF['01Agc'],DF['01机组出力'],DF['01储能']+DF['02储能']
    data['Pall'] = data['Pdg']+data['Pbat']
'''
def makedir(file):
    folder = os.path.exists(file)
    if not folder:
        os.makedirs(file)
        print("--- new folder ---")
        print("--- OK ---")
    else:
        print("--- the folder is exist ---")

def logit(logfile='out.log'):
    def logging_decorator(func):
        @wraps(func)
        def wrapped_function(*args,**kwargs):
            log_string = datetime.datetime.today().strftime('%Y-%m-%d %H:%M:%S')+':'+func.__name__ +" was called"
            with open(logfile,'a') as opened_file:
                opened_file.write(log_string + '\n')
            return func(*args,**kwargs)
        return wrapped_function
    return logging_decorator

# logit是自己编写的装饰器，现在用它装饰main_content_text_update函数

@logit(logfile='C:\\Users\\JesisW\\Desktop\\log.txt')
def main_content_text_update(doc,old,new,style=None):
    """其中替换的文字必须为同一组run，否则刚好在不同组run中将无法识别"""
    """在不同组run中极可能是由于各种样式、格式或中英文字体导致的"""
    """遇到上述情况，要么更改保持一致，要么采用段落文字更改"""
    for para in doc.paragraphs:
        for run in para.runs:
            if old in run.text:
                style = para.style
                run.text = run.text.replace(old,new)
                para.style = style
                if style is not None:
                    para.style = style

@logit(logfile='C:\\Users\\JesisW\\Desktop\\log.txt')
def tabel_text_update(doc,old,new,p=True,k=5):
    if p:
        '单纯进行表格内容替换'
        for tbl in doc.tables:
            for cell in tbl._cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old in run.text:
                            style = para.style
                            run.text = run.text.replace(old,new)
                            para.style = style
    else:
        '数据表格替换'
        m=k
        for tbl in doc.tables:
            for cell in tbl._cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if m>k-2:
                            if old in run.text:
                                m = 0
                                style = para.style
                                run.text = run.text.replace(old,new[m])
                                para.style = style
                                
                        else:
                            m += 1
                            style = para.style
                            run.text = str(new[m])
                            para.style = style

@logit(logfile='C:\\Users\\JesisW\\Desktop\\log.txt')
def pic_update(doc,old,pic_path):
    """图片插入，必须在此行段落进行插入，doc.add_picture是加在文末"""
    for para in doc.paragraphs:
        for run in para.runs:
            if old in run.text:
                style = para.style
                run.text = run.text.replace(old,'')
                pic = para.add_run().add_picture(pic_path,width=Inches(5))
                para.style = style

@logit(logfile='C:\\Users\\JesisW\\Desktop\\log.txt')
def header_update(doc,old,new):
    """对页眉页脚进行更改，在section部分"""
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                if old in run.text:
                    style = para.style
                    run.text = run.text.replace(old,new)
                    para.style = style

@logit(logfile='C:\\Users\\JesisW\\Desktop\\log.txt')
def toc_updates(doc):
    """目录的自动更新方法，但是此方法仍旧有局限：需要在打开文档时选择更改文档作用域才可"""
    name_space = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    update_name_space = "%supdateFields" % name_space
    val_name_space = "%sval" % name_space
    try:
        element_update_filed_obj = lxml.etree.SubElement(doc.settings.element,update_name_space)
        element_update_filed_obj.set(val_name_space,'true')
    except Exception as e:
        del e

def report_operation(Result_data,picpath,Aim_names_Chinese,sta_data_date):
    """运营报告配置信息"""
    datetoday = datetime.date.today()
    str_datetoday = datetoday.strftime('%Y.%m.%d')
    dateaskday = datetoday - datetime.timedelta(days=1)
    str_dateaskday = dateaskday.strftime('%Y.%m.%d')
    headers = {'station':Aim_names_Chinese,'year':str_datetoday}
    first_sheet_table = {'station':Aim_names_Chinese,'author':'王宝源','datetime':str_datetoday}
    toc_sheet = {'station':Aim_names_Chinese,'askdate':str_dateaskday,'askpeople':'毛一洋',
                     'reportdate':str_datetoday,'reportpeople':'王宝源','briefing':''}
    toc_table = {}
    main_content = {'station':Aim_names_Chinese,'number':'2'}
    main_content_kp = {'date1':'','date2':'','date3':'','date4':'','date5':''}
    main_content_CR = {'date6':'','date7':'','date8':'','date9':'','datex':''}
    main_content_pic_k = {'pic1':'','pic2':'','pic3':'','pic4':''}
    main_content_pic_CB = {'pic5':'','pic6':'','pic7':'','pic8':''}
    main_content_conclusion = {'conclusion1':'','conclusion2':''}
    
    document = Document('E:\\1伪D盘\\报告文档\\运营报告\\运营报告—模板.docx')
    df = Result_data[['time','k1','k2','k3','kp','D','Revenue','Cost','elecFee']]
    df.iloc[:,1:9] = df.iloc[:,1:9].round(2)
    df['time'] = df['time'].apply(lambda x:x.strftime('%Y/%m/%d'))
    k1,k2,k3,kp = df.iloc[-1,1],df.iloc[-1,2],df.iloc[-1,3],df.iloc[-1,4]
    i = -6
    for key in main_content_kp:
        i += 1
        L = df.iloc[i,0:6].to_list()
        main_content_kp[key] = L
    i = -6
    for key in main_content_CR:
        i += 1
        L = df.iloc[i,[0,6,7,8]].to_list()
        L.insert(1,round(L[1]-L[2]-L[3],2))
        L.insert(2,round(L[1]/L[2],2))
        main_content_CR[key] = L
    df2 = df.iloc[-5:,0:9]
    df2.loc[:,9]=df2.iloc[:,6]-df2.iloc[:,7]-df2.iloc[:,8]
    df2.loc[:,10] = df2.iloc[:,9]/df2.iloc[:,6]
    kpmax = df2['kp'].max()
    kpstd = df2['kp'].std()
    Revenuemax = df2.iloc[:,9].max()
    Revenuemaxrate = df2.iloc[:,10].max()
    if kpstd<0.075:
        word1 = 'kp变化较为平稳，没有明显的波动'
    else:
        word1 = 'kp变化较大'
    main_content_conclusion['conclusion1'] = '近5日内，kp最大值为%.2f，' %kpmax +word1
    main_content_conclusion['conclusion2'] = '近5日内，最大收益为%.2f元，最大收益率为%.2f' %(Revenuemax,Revenuemaxrate)
    pic_name = ['K1.png','K2.png','K3.png','Kp.png','收益图.png','成本图.png','收益成本图.png','电费图.png']
    i = 0
    for key in main_content_pic_k:
        main_content_pic_k[key]=picpath+'\\'+pic_name[i]
        i += 1
    toc_sheet['briefing'] = '本日仿真k1为%.2f，k2为%.2f，k3为%.2f，kp为%.2f\n' %(k1,k2,k3,kp)+\
    main_content_conclusion['conclusion1']+'\n'+main_content_conclusion['conclusion2']
    for old,new in headers.items():
        header_update(document, old, new)
    
    for old,new in first_sheet_table.items():
        tabel_text_update(document, old, new, p=True)
    
    for old,new in toc_sheet.items():
        tabel_text_update(document, old, new, p=True)
    
    for old,new in main_content.items():
        main_content_text_update(document, old, new)
    
    for old,new in main_content_kp.items():
        tabel_text_update(document, old, new, p=False,k=len(main_content_kp[old]))
    
    for old,new in main_content_CR.items():
        tabel_text_update(document, old, new, p=False,k=len(main_content_CR[old]))
    
    for old,pic_path in main_content_pic_k.items():
        pic_update(document, old, pic_path)
        
    for old,pic_path in main_content_pic_CB.items():
        pic_update(document, old, pic_path)
    
    for old,new in main_content_conclusion.items():
        main_content_text_update(document, old, new)
    
    toc_updates(document)
    
    word_name = 'E:\\1伪D盘\\报告文档\\运营报告\\'+Aim_names_Chinese+'\\运营报告'+sta_data_date[0:10]+Aim_names_Chinese+'.docx'
    pdf_name = 'E:\\1伪D盘\\报告文档\\运营报告\\'+Aim_names_Chinese+'\\运营报告'+sta_data_date[0:10]+Aim_names_Chinese+'.pdf'
    document.save(word_name)
    word2pdf(word_name, pdf_name)
    return 

def report_running(Result_data,Agcresult,Batresult,pic_path,Aim_names_Chinese,sta_data_date,Jizu):
    """运行报告配置信息"""
    document = Document('E:\\1伪D盘\\报告文档\\运行报告\\运行报告—模板.docx')
    Stationname = ['新丰','云河','海丰','河源','鲤鱼江','恒运','宣化','准大','兴和','上都','平朔','同达']
    Area =        ['蒙西','广东','广东','广东','广东' ,'广东','华北','蒙西','蒙西','华北','华北','华北']
    BatPe = [9,9,30,18,12,15,9,9,9,18,9,9] # 电池的额定功率，固定值，写死
    BatPe = dict(zip(Stationname,BatPe))
    AREA = dict(zip(Stationname,Area))
    Pe = [300,330,1050,600,300,300,330,300,300,600,300,300] # 机组额定功率
    Pe = dict(zip(Stationname,Pe))
    P = Pe[Aim_names_Chinese]
    if AREA[Aim_names_Chinese] == '蒙西':
        Vn = P*0.015
        konemax = 2
    elif AREA[Aim_names_Chinese] == '广东':
        Vn = P*0.01751#0.01903
        konemax = 5
    elif AREA[Aim_names_Chinese] == '华北':
        Vn = P*0.015
        konemax = 1.5
    Vn = round(Vn,2)
    Vntwo = 2*Vn
    
    datetoday = datetime.date.today()
    str_datetoday = datetoday.strftime('%Y.%m.%d')
    dateaskday = datetoday - datetime.timedelta(days=1)
    str_dateaskday = dateaskday.strftime('%Y.%m.%d')
    headers = {'station':Aim_names_Chinese,'year':str_datetoday}
    first_sheet_table = {'station':Aim_names_Chinese,'author':'王宝源','datetime':str_datetoday}
    toc_sheet = {'station':Aim_names_Chinese,'askdate':str_dateaskday,'askpeople':'毛一洋',
                     'reportdate':str_datetoday,'reportpeople':'王宝源','briefing':''}
    toc_table = {}
    main_content = {'Vb':str(Vn),'Vc':str(Vntwo),'konemax':str(konemax)}
    main_content_pic = {'pic1':'','pic2':'','pic3':'','pic4':'','pic5':'','pic6':'','pic7':''}
    main_content_cyc = {'date1':'','date2':'','date3':'','date4':'','date5':''}
    main_content_conclusion = {'conclusion1':'','conclusion2':'','conclusion3':'','conclusion4':''}
    Con1_choose = '本日AGC指令分析发现， AGC指令调节主要以%s调节为主，AGC指令调节程度%s，调节速率%s，平均调节时间%s，折返调节%s。'
    Con2_choose = '对于储能的运行分析部主要以放电深度和等效放电时长作为手段进行分析，通过分析发现，本日储能%s响应AGC指令(不排除数据源存在异常)，单次最大充电%.2f%%，单次最大放电为%.2f%%。'
    
    Con3_choose = '今日等效循环次数约为%.2f，充电比放电多%.2f个循环，充电、放电比为%.2f:1，充电比放电多%.2f MWh。'
    
    Con4_choose = '根据上述原则，计算得到%s站%s号机组的机组调节调节分析：\n\
    1a）不调节的比例为%.2f%%；\n\
    1b）调节缓慢的比例为%.2f%%；\n\
    2）反调的比例为%.2f%%；\n\
    3）瞬间调节完成的比例为%.2f%%；\n\
    4）在问题调节中，储能有效调节的比例为%.2f%%。\n'
    main_content_conclusion = {'conclusion1':Con1_choose,'conclusion2':Con2_choose,'conclusion3':Con3_choose,'conclusion4':Con4_choose}
    pic_name = ['蜡烛图.png','调节速度.png','折返.png','平均调节时间图.png','DOD.png','等效2C时长图.png','等效循环图.png']
    i = 0
    for key in main_content_pic:
        main_content_pic[key] = pic_path+'\\'+pic_name[i]
        i += 1
    df = Result_data[['time','充电等效次数','放电等效次数']]
    df.loc[:,'time'] = df.loc[:,'time'].apply(lambda x:x.strftime('%Y/%m/%d'))
    word7 = df.iloc[-1,2]
    word8 = abs(df.iloc[-1,1])-df.iloc[-1,2]
    word9 = abs(df.iloc[-1,1])/df.iloc[-1,2]
    df.loc[:,'充电量'] = df.loc[:,'充电等效次数']*BatPe[Aim_names_Chinese]
    df.loc[:,'放电量'] = df.loc[:,'放电等效次数']*BatPe[Aim_names_Chinese]
    df = df.iloc[:,:].round(2)
    word10 = abs(df['充电量'][-1])-df['放电量'][-1] 
    i = -6
    for key in main_content_cyc:
        i += 1
        main_content_cyc[key] = df.iloc[i,:].to_list()
    df = Agcresult[['Max','Min','Beg','End']]
    df.loc[:,'Std'] = df.std(axis=1)
    a = len(df[df['Std']>10])/96
    b = 1-a
    if a/b>4:
        word1 = '折返调节'
        word2 = '频繁'
    elif a/b<0.25:
        word1 = '不变负荷调节'
        word2 = '稀疏'
    else:
        word1 = '折返调节和不变负荷调节'
        word2 = '适中'
    df = Agcresult['SaD']
    a = len(df[df>Vn])/96
    if a>0.67:
        word3 = '要求较高'
    elif a>0.33:
        word3 = '要求适中'
    else:
        word3 = '要求较低'
    df = Agcresult['Avet']
    a = len(df[df<100])/96
    if a>0.67:
        word4 = '较短'
    elif a>0.33:
        word4 = '适中'
    else:
        word4 = '较长'
    df = Agcresult[['Ret','Num']]
    a = df['Num']==0
    df['Num'][a] = 1
    df.loc[:,'proption'] = df.loc[:,'Ret']/df.loc[:,'Num']
    a = len(df[df['proption']>0.4])/96
    if a >0.5:
        word5 = '严重'
    elif a>0.3:
        word5 = '适中'
    else:
        word5 = '较少'    
    Con1_choose = Con1_choose %(word1,word2,word3,word4,word5)
    
    df = Batresult['DOD']
    a = len(df[df>5])/96
    if a >0.7:
        word6 = '频繁响应Agc'
    elif a>0.3:
        word6 = '部分响应Agc'
    else:
        word6 = '较少响应Agc'
    a = df.max()
    b = df.min()
    Con2_choose = Con2_choose %(word6,b,a)
    Con3_choose = Con3_choose %(word7,word8,word9,word10)
    df = Result_data.loc[Result_data.index[-1],['不动','缓调','反调','瞬间完成','比例']]
    Con4_choose = Con4_choose %(Aim_names_Chinese,Jizu,df['不动'],df['缓调'],df['反调'],df['瞬间完成'],df['比例'])
    toc_sheet['briefing'] = Con1_choose+'\n'+Con3_choose
    main_content_conclusion = {'conclusion1':Con1_choose,'conclusion2':Con2_choose,'conclusion3':Con3_choose,'conclusion4':Con4_choose}
    
    for old,new in headers.items():
        header_update(document, old, new)
    
    for old,new in first_sheet_table.items():
        tabel_text_update(document, old, new, p=True)
    
    for old,new in toc_sheet.items():
        tabel_text_update(document, old, new, p=True)
    
    for old,new in main_content.items():
        main_content_text_update(document, old, new)
    
    for old,new in main_content_cyc.items():
        tabel_text_update(document, old, new, p=False,k=len(main_content_cyc[old]))
    
    for old,pic_path in main_content_pic.items():
        pic_update(document, old, pic_path)
    
    for old,new in main_content_conclusion.items():
        main_content_text_update(document, old, new)
    
    toc_updates(document)
    
    word_name = 'E:\\1伪D盘\\报告文档\\运行报告\\'+Aim_names_Chinese+'\\运行报告'+sta_data_date[0:10]+Aim_names_Chinese+'.docx'
    pdf_name = 'E:\\1伪D盘\\报告文档\\运行报告\\'+Aim_names_Chinese+'\\运行报告'+sta_data_date[0:10]+Aim_names_Chinese+'.pdf'
    document.save(word_name)
    word2pdf(word_name, pdf_name)
    return

def word2pdf(word_name,pdf_name):
    try:
        word = client.DispatchEx('Word Application')
        worddoc = word.Documents.Open(word_name,ReadOnly = 1)
        worddoc.SaveAs(pdf_name,FileFormat = 17)
        worddoc.close()
        return 
    except:
        print('打印失败')
        return 

if __name__ == '__main__':
    sta_config=pd.read_excel('D:/pyworkspace/Data_Get/assets/sta_cl_ah_config.xlsx')
    index =      [  0,    1,    2,   5,    6,    7,   14,   15,   16,  33,  4]  #调频电站序号
    sta_data_date = '2019-12-16 00:00:00'
    end_data_date = '2019-12-17 00:00:00'
    dianzhan      =        3
    jizu = input('请输入机组：')
    jizu = int(jizu)
    if jizu == 1:
        Jizu = '01'
    else:
        Jizu = '02'

    if jizu not in [1,2]:
        print('错误输入机组参数，请重新输入')
        os._exit(0) # 在子程序中就是退出子程序
#     for i in sta_config['name']:
#         file = r'E:\1伪D盘\AGC运行'+'\\'+i+'\\图片'
#         makedir(file)
    # 自动获取程序
    path = r'E:\1伪D盘\AGC运行'
#     K = Check_data(dianzhan)
#     Aim_names,Aim_names_Chinese = K.Aim_names,K.Aim_names_Chinese
#     index_loc = K.location
#     data,time,Bat,Agc = K.webfile(jizu, sta_data_date, end_data_date)
#     if data.empty:
#         os._exit(0) # 在子程序中就是退出子程序
#     Bat.to_csv(path+'\\'+Aim_names+'\\数据\\储能Data'+sta_data_date[0:10]+'-'+str(jizu)+'.csv',index = True,header=True,encoding='gbk')
#     Agc.to_csv(path+'\\'+Aim_names+'\\数据\\AGCData'+sta_data_date[0:10]+'-'+str(jizu)+'.csv',index = True,header=True,encoding='gbk')
#     Bat = Bat['01储能']+Bat['02储能']
#     Pdg = Agc[Jizu+'机组出力']
#     Agc = Agc[Jizu+'AGC']

#     人工操作
    Aim_names_Chinese = '海丰'
    index_loc = 'GD'
    a = pd.read_csv(r'E:\1伪D盘\AGC运行\广东华润海丰电厂储能电站\数据\储能Data2019-12-16-2.csv',index_col = 0,header = 0,encoding='gbk')
    b = pd.read_csv(r'E:\1伪D盘\AGC运行\广东华润海丰电厂储能电站\数据\AGCData2019-12-16-2.csv',index_col = 0,header = 0,encoding='gbk')
    df = pd.merge(a,b,left_index=True,right_index=True,how='inner')
    data = pd.DataFrame(columns=['Agc','Pdg','Pall','Pbat'])
    data['Agc'],data['Pdg'],data['Pbat'] = df[Jizu+'AGC'],df[Jizu+'机组出力'],df['01储能']+df['02储能']
    data['Pall'] = data['Pdg']+data['Pbat']
    Bat = a['01储能']+a['02储能']
    Agc = b[Jizu+'AGC']
    Pdg = b[Jizu+'机组出力']
    time = pd.to_datetime(df.index,format='%Y-%m-%d %H:%M:%S')
    time.columns = ['time']
    Aim_names = sta_config['name'][index[dianzhan]]

#     df = pd.read_csv(r'D:\华润海丰\Data\海丰项目26-29日AGC指令数据\29号.csv',encoding ='gbk',index_col=0,header=0)
#     df.index = df['Date']
#     time = pd.to_datetime(df['Date'],format='%Y/%m/%d %H:%M:%S')
#     time.columns = ['time']
#     data = pd.DataFrame(columns=['Agc','Pdg','Pall','Pbat'])
#     data['Agc'],data['Pdg'],data['Pall'],data['Pbat'] = df['01AGC'],df['01机组出力'],df['Pall'],df['bat']
#     Bat = df['bat']
#     Agc = df['01AGC']
#     Pdg = df['01机组出力']
#     Aim_names_Chinese = '海丰'
#     index_loc = 'GD'
#     Aim_names = sta_config['name'][index[dianzhan]]
    # 处理数据
    file = r'E:\1伪D盘\AGC运行'+'\\'+Aim_names+'\\图片'+'\\'+sta_data_date[0:10]
    makedir(file)
    deal = deal_data(Aim_names_Chinese=Aim_names_Chinese,data=data,location=index_loc,Agc=Agc,Pdg=Pdg,Bat=Bat,time=time)
    Result,Result1,BatResult,ResultSingle = deal.Solv()
#     kp,k1,k2,k3,D,Result_Agc = deal.Solv()
    # 保存数据
    ResultSingle['机组'] = jizu
    ResultSingle.index = [sta_data_date[0:10]]
    ResultSingle.index.name = 'time'
    ResultSingle.index = pd.to_datetime(ResultSingle.index,format=None)
    if os.path.exists(path+'\\'+Aim_names+'\\Kp\\data.csv'):
        ResultSingle1 = pd.read_csv(path+'\\'+Aim_names+'\\Kp\\data.csv',index_col=0,header=0,encoding='gbk')
        ResultSingle1.index = pd.to_datetime(ResultSingle1.index,format=None)
        ResultSingle1.to_csv(path+'\\'+Aim_names+'\\Kp\\备份data.csv',index = True,header=True,encoding='gbk')
#         # 手工画历史kp图
#         fig2 = FIGUREplot.figecharts(i=dianzhan,kpresult=ResultSingle1)
#         fig2.figKp(file)
        if ResultSingle.index[0] not in ResultSingle1.index:
            ResultSingle = pd.concat([ResultSingle1,ResultSingle],axis = 0)
            ResultSingle.sort_index(axis =0,inplace=True)
            ResultSingle.to_csv(path+'\\'+Aim_names+'\\Kp\\data.csv',index = True,header=True,encoding='gbk')
        else:
            ResultSingle = ResultSingle1
            print('kp重复计算，请检查')
    else:
        ResultSingle.to_csv(path+'\\'+Aim_names+'\\Kp\\data.csv',index = True,header=True,encoding='gbk')
    Result.to_csv(path+'\\'+Aim_names+'\\AGC'+'\\'+sta_data_date[0:10]+Aim_names_Chinese+'-'+str(jizu)+'.csv',index=True,header=True,encoding='gbk')
    Result1.to_csv(path+'\\'+Aim_names+'\\机组'+'\\'+sta_data_date[0:10]+Aim_names_Chinese+'-'+str(jizu)+'.csv',index=True,header=True,encoding='gbk')
    BatResult.to_csv(path+'\\'+Aim_names+'\\储能'+'\\'+sta_data_date[0:10]+Aim_names_Chinese+'-'+str(jizu)+'.csv',index=True,header=True,encoding='gbk')
#     # 画其他图
#     fig1 = FIGUREplot.figecharts(i=dianzhan,Agcresult=Result,BatResult=BatResult)
#     fig1.figAgc(file)
#     fig1.figBat(file)
#     # 画kp图-日更新的
#     fig2 = FIGUREplot.figecharts(i=dianzhan,kpresult=ResultSingle)
#     fig2.figKp(file)
#     ResultSingle = pd.read_csv(r'E:\1伪D盘\AGC运行\广东华润海丰电厂储能电站\Kp\data.csv',header=0,encoding='gbk')
#     Result = pd.read_csv(r'E:\1伪D盘\AGC运行\广东华润海丰电厂储能电站\AGC\2019-12-16海丰-2.csv',header=0,encoding='gbk')
#     BatResult = pd.read_csv(r'E:\1伪D盘\AGC运行\广东华润海丰电厂储能电站\储能\2019-12-16海丰-2.csv',header=0,encoding='gbk')
    ResultSingle['time'] = ResultSingle.index
    report_operation(ResultSingle, file, Aim_names_Chinese, sta_data_date)
    report_running(ResultSingle, Result, BatResult, file, Aim_names_Chinese, sta_data_date,Jizu)